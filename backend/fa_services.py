"""
fa_services.py — Core service layer for the FastAPI backend.

Direct port of the original Django services.py with Django ORM calls
replaced by SQLAlchemy session calls.  All LangChain logic is unchanged.
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Generator

import numpy as np

# LangChain Imports (identical to Django version)
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_groq import ChatGroq
from langchain_community.chat_models import ChatOllama

import config

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# EmbeddingService
# ---------------------------------------------------------------------------

class EmbeddingService:
    _embeddings = None

    @classmethod
    def get_embeddings(cls):
        if cls._embeddings is None:
            logger.info("Loading embedding model: %s", config.EMBEDDING_MODEL_NAME)
            cls._embeddings = HuggingFaceEmbeddings(
                model_name=config.EMBEDDING_MODEL_NAME,
                encode_kwargs={"normalize_embeddings": True},
            )
        return cls._embeddings

    @classmethod
    def encode(cls, text: str) -> List[float]:
        return cls.get_embeddings().embed_query(text)

    @classmethod
    def encode_batch(cls, texts: List[str]) -> List[List[float]]:
        return cls.get_embeddings().embed_documents(texts)


# ---------------------------------------------------------------------------
# SemanticChunker — structure-aware chunking
# ---------------------------------------------------------------------------

class _RawChunk:
    """Intermediate chunk before embedding."""
    __slots__ = ("text", "page_number", "chunk_type", "section_heading", "meta")

    def __init__(self, text: str, page_number: int, chunk_type: str,
                 section_heading: str = "", meta: dict = None):
        self.text = text.strip()
        self.page_number = page_number
        self.chunk_type = chunk_type          # heading | paragraph | table | image | caption
        self.section_heading = section_heading
        self.meta = meta or {}               # rich structural metadata


class SemanticChunker:
    """
    Splits documents on structural boundaries:
      PDF  → PyMuPDF block-level analysis (text blocks classified as heading /
             paragraph / image; table regions detected via find_tables())
      DOCX → python-docx paragraph styles (Heading 1/2/3, Normal, Table)
      TXT  → double-newline paragraph splits
    Small consecutive paragraphs are merged up to MAX_CHARS to avoid
    fragmentation.  Heading text is prepended to the next paragraph as context.
    """

    MAX_CHARS    = 900   # soft cap for a single chunk
    MIN_CHARS    = 60    # minimum — ignore noise / whitespace-only blocks

    # ------------------------------------------------------------------ PDF
    @classmethod
    def from_pdf(cls, file_path: str) -> List[_RawChunk]:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        page_width  = doc[0].rect.width  if len(doc) else 0
        page_height = doc[0].rect.height if len(doc) else 0
        raw: List[_RawChunk] = []

        for page_num, page in enumerate(doc, start=1):
            # --- collect table rectangles so we can skip those text blocks ---
            table_rects = []
            try:
                tables = page.find_tables()
                for tbl in tables.tables:
                    table_rects.append(tbl.bbox)
                    rows_data = tbl.extract()
                    n_rows = len(rows_data)
                    n_cols = len(rows_data[0]) if rows_data else 0
                    rows = []
                    for row in rows_data:
                        cells = [str(c or "").strip() for c in row]
                        rows.append(" | ".join(cells))
                    table_text = "\n".join(rows)
                    if len(table_text) >= cls.MIN_CHARS:
                        raw.append(_RawChunk(
                            table_text, page_num, "table",
                            meta={
                                "rows": n_rows,
                                "cols": n_cols,
                                "bbox": list(tbl.bbox),
                                "page_width":  round(page_width, 1),
                                "page_height": round(page_height, 1),
                            }
                        ))
            except Exception:
                pass  # PyMuPDF versions without find_tables()

            # --- page font stats for heading detection ---
            blocks = page.get_text("dict", flags=11)["blocks"]
            font_sizes = []
            for b in blocks:
                if b["type"] != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        font_sizes.append(span["size"])

            avg_size = (sum(font_sizes) / len(font_sizes)) if font_sizes else 12.0
            heading_threshold = avg_size * 1.15   # 15 % bigger = heading candidate

            # --- process blocks ---
            for b in blocks:
                if b["type"] == 1:
                    # Image block
                    raw.append(_RawChunk(
                        "[Image]", page_num, "image",
                        meta={
                            "bbox": list(b["bbox"]),
                            "page_width":  round(page_width, 1),
                            "page_height": round(page_height, 1),
                        }
                    ))
                    continue

                if b["type"] != 0:
                    continue

                # Check whether this text block is inside a known table rect
                bx0, by0, bx1, by1 = b["bbox"]
                in_table = any(
                    tx0 <= bx0 and ty0 <= by0 and bx1 <= tx1 and by1 <= ty1
                    for (tx0, ty0, tx1, ty1) in table_rects
                )
                if in_table:
                    continue

                # Assemble text + detect type
                lines_text = []
                is_bold_block = False
                max_size = 0.0
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        lines_text.append(span["text"])
                        if span["size"] > max_size:
                            max_size = span["size"]
                        if span["flags"] & 16:   # bold flag
                            is_bold_block = True

                block_text = " ".join(lines_text).replace("\x00", "").strip()
                if not block_text or len(block_text) < cls.MIN_CHARS:
                    continue

                # Classify
                is_heading = (max_size >= heading_threshold) or is_bold_block
                chunk_type = "heading" if is_heading else "paragraph"
                raw.append(_RawChunk(
                    block_text, page_num, chunk_type,
                    meta={
                        "bbox":          [round(v, 1) for v in b["bbox"]],
                        "font_size":     round(max_size, 1),
                        "avg_font_size": round(avg_size, 1),
                        "is_bold":       is_bold_block,
                        "page_width":    round(page_width, 1),
                        "page_height":   round(page_height, 1),
                    }
                ))

        doc.close()
        return cls._merge(raw)

    # ------------------------------------------------------------------ DOCX
    @classmethod
    def from_docx(cls, file_path: str) -> List[_RawChunk]:
        import docx as _docx

        document = _docx.Document(file_path)
        raw: List[_RawChunk] = []

        for para in document.paragraphs:
            text = para.text.replace("\x00", "").strip()
            if not text or len(text) < cls.MIN_CHARS:
                continue
            style_name = para.style.name or ""
            style_lower = style_name.lower()
            if "heading" in style_lower:
                chunk_type = "heading"
                # extract heading level (e.g. "Heading 2" → 2)
                try:
                    heading_level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1
            elif "caption" in style_lower:
                chunk_type = "caption"
                heading_level = 0
            else:
                chunk_type = "paragraph"
                heading_level = 0
            raw.append(_RawChunk(
                text, 1, chunk_type,
                meta={
                    "style":         style_name,
                    "heading_level": heading_level,
                }
            ))

        for tbl_idx, tbl in enumerate(document.tables):
            rows_data = []
            for row in tbl.rows:
                cells = [c.text.replace("\x00", "").strip() for c in row.cells]
                rows_data.append(" | ".join(cells))
            table_text = "\n".join(rows_data)
            if len(table_text) >= cls.MIN_CHARS:
                raw.append(_RawChunk(
                    table_text, 1, "table",
                    meta={
                        "rows":      len(tbl.rows),
                        "cols":      len(tbl.columns),
                        "table_idx": tbl_idx,
                    }
                ))

        return cls._merge(raw)

    # ------------------------------------------------------------------ TXT
    @classmethod
    def from_txt(cls, file_path: str) -> List[_RawChunk]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read().replace("\x00", "")
        paragraphs = [p.strip() for p in content.split("\n\n") if len(p.strip()) >= cls.MIN_CHARS]
        raw = [_RawChunk(p, 1, "paragraph") for p in paragraphs]
        return cls._merge(raw)

    # ------------------------------------------------------------------ MERGE
    @classmethod
    def _merge(cls, raw: List[_RawChunk]) -> List[_RawChunk]:
        """
        Merge consecutive small paragraphs; prepend heading context.
        Headings, tables, and images are always kept as standalone chunks.
        Metadata from the first block in a merged group is preserved.
        """
        merged: List[_RawChunk] = []
        current_heading = ""
        buffer_text = ""
        buffer_page = 1
        buffer_meta: dict = {}

        def flush(page):
            nonlocal buffer_text, buffer_meta
            if buffer_text.strip():
                text = f"[{current_heading}]\n{buffer_text}" if current_heading else buffer_text
                meta = {**buffer_meta, "section": current_heading}
                merged.append(_RawChunk(text, page, "paragraph", current_heading, meta))
            buffer_text = ""
            buffer_meta = {}

        for chunk in raw:
            if chunk.chunk_type in ("table", "image"):
                flush(chunk.page_number)
                text = f"[{current_heading}]\n{chunk.text}" if current_heading and chunk.chunk_type == "table" else chunk.text
                meta = {**chunk.meta, "section": current_heading}
                merged.append(_RawChunk(text, chunk.page_number, chunk.chunk_type, current_heading, meta))

            elif chunk.chunk_type == "heading":
                flush(chunk.page_number)
                current_heading = chunk.text
                meta = {**chunk.meta, "section": current_heading}
                merged.append(_RawChunk(chunk.text, chunk.page_number, "heading", "", meta))

            else:  # paragraph / caption
                if len(buffer_text) + len(chunk.text) > cls.MAX_CHARS:
                    flush(chunk.page_number)
                if not buffer_text:
                    buffer_meta = chunk.meta  # take meta from first block in group
                buffer_text += (" " if buffer_text else "") + chunk.text
                buffer_page = chunk.page_number

        flush(buffer_page)
        return merged


# ---------------------------------------------------------------------------
# IngestionService
# ---------------------------------------------------------------------------

class IngestionService:
    @staticmethod
    def ingest(db, document) -> int:
        """
        Structure-aware ingestion pipeline.
        If the document already has chunks, they are REPLACED (delete + re-insert).
        Returns the number of chunks created.
        """
        from db_models import Chunk

        # ── 0. Delete existing chunks for this document (replace, don't duplicate)
        existing = db.query(Chunk).filter(Chunk.document_id == document.id).count()
        if existing:
            logger.info("Replacing %d existing chunks for document %d", existing, document.id)
            db.query(Chunk).filter(Chunk.document_id == document.id).delete()
            db.commit()

        # ── 1. Semantic extraction
        file_path = str(config.MEDIA_ROOT / document.file)
        file_type = document.file_type.lower().strip(".")

        try:
            if file_type == "pdf":
                raw_chunks = SemanticChunker.from_pdf(file_path)
            elif file_type == "docx":
                raw_chunks = SemanticChunker.from_docx(file_path)
            elif file_type == "txt":
                raw_chunks = SemanticChunker.from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error("Semantic chunking failed for %s: %s", document.title, e)
            return 0

        if not raw_chunks:
            return 0

        # Filter empty/too-short
        raw_chunks = [c for c in raw_chunks if len(c.text) >= SemanticChunker.MIN_CHARS]
        if not raw_chunks:
            return 0

        # ── 2. Embed (batch)
        texts = [c.text for c in raw_chunks]
        embeddings = EmbeddingService.encode_batch(texts)

        # ── 3. Persist
        max_gi = db.query(Chunk.global_index).order_by(Chunk.global_index.desc()).first()
        current_gi = ((max_gi[0] or -1) + 1) if max_gi else 0

        chunk_objs = []
        for i, (rc, emb) in enumerate(zip(raw_chunks, embeddings)):
            chunk_objs.append(
                Chunk(
                    document_id  = document.id,
                    content      = rc.text,
                    embedding    = emb,
                    page_number  = rc.page_number,
                    chunk_index  = i,
                    global_index = current_gi,
                    chunk_type   = rc.chunk_type,
                    chunk_meta   = rc.meta,
                )
            )
            current_gi += 1

        db.bulk_save_objects(chunk_objs)

        # Update total_pages
        max_page = max((c.page_number for c in chunk_objs), default=0)
        document.total_pages = max_page
        db.commit()

        logger.info(
            "Ingested %d semantic chunks for '%s' (%s)",
            len(chunk_objs), document.title, file_type
        )
        return len(chunk_objs)


# ---------------------------------------------------------------------------
# CustomRetriever
# ---------------------------------------------------------------------------

class CustomRetriever:
    @staticmethod
    def retrieve_documents(
        db,
        query: str,
        document_ids: List[int],
        top_k: int = 4,
    ) -> List[Dict]:
        from db_models import Chunk, Document

        if not document_ids:
            return []

        query_vec = np.array(EmbeddingService.encode(query))
        norm = np.linalg.norm(query_vec)
        if norm > 0:
            query_vec = query_vec / norm

        results = []
        for doc_id in document_ids:
            rows = (
                db.query(
                    Chunk.id,
                    Chunk.document_id,
                    Document.title.label("document_title"),
                    Chunk.content,
                    Chunk.page_number,
                    Chunk.embedding,
                    Chunk.global_index,
                    Chunk.chunk_type,
                    Chunk.chunk_meta,
                )
                .join(Document, Chunk.document_id == Document.id)
                .filter(Chunk.document_id == doc_id)
                .all()
            )
            if not rows:
                continue

            embs = np.array([r.embedding for r in rows])
            norms = np.linalg.norm(embs, axis=1, keepdims=True)
            norms[norms == 0] = 1.0
            embs_norm = embs / norms
            scores = np.dot(embs_norm, query_vec)

            scored = []
            for i, r in enumerate(rows):
                scored.append(
                    {
                        "chunk_id":       r.id,
                        "document_id":    r.document_id,
                        "document_title": r.document_title,
                        "content":        r.content,
                        "page_number":    r.page_number,
                        "global_index":   r.global_index,
                        "chunk_type":     getattr(r, "chunk_type", "paragraph"),
                        "chunk_meta":     getattr(r, "chunk_meta", {}) or {},
                        "score":          float(scores[i]),
                    }
                )
            scored.sort(key=lambda x: x["score"], reverse=True)
            results.extend(scored[:top_k])

        results.sort(key=lambda x: x["score"], reverse=True)
        return results

    @staticmethod
    def to_langchain_docs(scored_chunks: List[Dict]) -> List[LCDocument]:
        docs = []
        for c in scored_chunks:
            docs.append(
                LCDocument(
                    page_content=c["content"],
                    metadata={
                        "score":          c.get("score"),
                        "document_title": c.get("document_title"),
                        "page_number":    c.get("page_number"),
                        "document_id":    c.get("document_id"),
                    },
                )
            )
        return docs


# ---------------------------------------------------------------------------
# RAGService
# ---------------------------------------------------------------------------

class RAGService:
    @staticmethod
    def get_llm():
        if config.GROQ_API_KEY:
            return ChatGroq(
                temperature=0.3,
                model_name=config.GROQ_MODEL,
                api_key=config.GROQ_API_KEY,
                streaming=True,
            )
        return ChatOllama(
            model=config.OLLAMA_MODEL,
            temperature=0.3,
            streaming=True,
        )

    @staticmethod
    def stream_response(
        query: str,
        docs: List[LCDocument],
        intent: str = "rag",
        history: List[dict] = None,
    ) -> Generator[str, None, None]:
        llm = RAGService.get_llm()

        # Build context grouped by document
        doc_map: Dict[str, List[LCDocument]] = {}
        for d in docs:
            title = d.metadata.get("document_title", "Unknown")
            doc_map.setdefault(title, []).append(d)

        context_parts = []
        for title, d_list in doc_map.items():
            context_parts.append(f"Document: {title}")
            for d in d_list:
                page = d.metadata.get("page_number", "?")
                context_parts.append(f"[p.{page}] {d.page_content}")
            context_parts.append("")
        context_str = "\n".join(context_parts)

        messages = []
        if intent == "mermaid":
            messages.append((
                "system",
                "You are a diagram generator. Output ONLY valid Mermaid.js syntax. "
                "Do NOT include explanation, markdown fences, or text. "
                "Start with the diagram type keyword (e.g. flowchart, sequenceDiagram).",
            ))
            if history:
                for h in history:
                    messages.append((h["role"], h["content"]))
            messages.append(("user", "Based on:\n{context}\n\nGenerate a Mermaid diagram for: {query}"))

        elif intent == "compare":
            messages.append((
                "system",
                "You are a helpful research analyst. Compare the provided documents clearly. "
                "Output Format:\n1. **Agreements**\n2. **Differences**\n3. **Supporting Evidence** (Cite [p. X])\n"
                "Do NOT use outside knowledge.",
            ))
            if history:
                for h in history:
                    messages.append((h["role"], h["content"]))
            messages.append(("user", "DOCUMENTS:\n{context}\n\nCOMPARE:\n{query}"))

        else:
            messages.append((
                "system",
                "You are a friendly and helpful research assistant. "
                "You are having a conversation with the user about their documents. "
                "Answer their questions using ONLY the provided context. "
                "If the context doesn't contain the answer, say 'I couldn't find that in the documents'. "
                "Cite your sources as [p. X] where appropriate. "
                "Be conversational but precise.",
            ))
            if history:
                for h in history:
                    messages.append((h["role"], h["content"]))
            messages.append(("user", "CONTEXT:\n{context}\n\nQUESTION:\n{query}"))

        prompt = ChatPromptTemplate.from_messages(messages)
        chain = prompt | llm | StrOutputParser()
        return chain.stream({"context": context_str, "query": query})


# ---------------------------------------------------------------------------
# IntentRouter
# ---------------------------------------------------------------------------

class IntentRouter:
    MERMAID_KEYWORDS = {"draw", "diagram", "flow", "flowchart", "architecture", "chart", "visualize", "graph"}
    COMPARE_KEYWORDS = {"compare", "contrast", "difference", "vs", "versus", "both", "agreement", "disagree"}

    @staticmethod
    def classify(query: str) -> str:
        words = set(re.findall(r"\w+", query.lower()))
        if words & IntentRouter.MERMAID_KEYWORDS:
            return "mermaid"
        if words & IntentRouter.COMPARE_KEYWORDS:
            return "compare"
        return "rag"


# ---------------------------------------------------------------------------
# TitleService
# ---------------------------------------------------------------------------

class TitleService:
    @staticmethod
    def generate(text: str) -> str:
        preview = " ".join(text.split()[:1000])
        llm = RAGService.get_llm()
        prompt = ChatPromptTemplate.from_messages([
            ("system", "Generate a short 3-6 word title. Output ONLY the title."),
            ("user", "{text}"),
        ])
        chain = prompt | llm | StrOutputParser()
        try:
            return chain.invoke({"text": preview}).strip()[:200]
        except Exception:
            return "Research Session"
